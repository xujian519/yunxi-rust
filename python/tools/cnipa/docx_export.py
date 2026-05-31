#!/usr/bin/env python3
"""国知局版式 DOCX 导出 — Markdown → 专利意见陈述书 Word 文件。

用法：
    python3 docx_export.py <input.md> <output.docx>

依赖：
    pip install python-docx

格式说明：
    遵循中国国家知识产权局 (CNIPA) 专利审查指南规定的意见陈述书格式。
    正文使用宋体 (SimSun)，标题使用黑体 (SimHei) / 粗体。
    纸张 A4，上下边距 2.54cm，左右边距 3.17cm（默认）。
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError as e:
    print(f"错误: 缺少 python-docx 库。请执行: pip install python-docx\n{str(e)}", file=sys.stderr)
    sys.exit(2)


CNIPA_FONT = "仿宋"
CNIPA_FONT_EN = "Times New Roman"
TITLE_FONT = "黑体"
TITLE_FONT_EN = "Times New Roman"


def set_run_font(run, cn_name=CNIPA_FONT, en_name=CNIPA_FONT_EN, size=12, bold=False):
    """设置 run 的字体（中英文）。"""
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = en_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), cn_name)
    rFonts.set(qn("w:ascii"), en_name)
    rFonts.set(qn("w:hAnsi"), en_name)


def add_cnipa_styled_paragraph(doc, text, cn_font=CNIPA_FONT, en_font=CNIPA_FONT_EN, size=12, bold=False, alignment=None, space_after=Pt(6)):
    """添加带有 CNIPA 字体格式的段落。"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, cn_font, en_font, size, bold)
    return p


def add_heading_styled(doc, text, level=1):
    """添加 CNIPA 格式的标题。"""
    # 标题级别映射
    sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5

    # 国知局版式：一、二、三 级标题使用黑体
    title_cn = TITLE_FONT
    title_en = TITLE_FONT_EN
    bold = level <= 2

    run = p.add_run(text)
    set_run_font(run, title_cn, title_en, sizes.get(level, 12), bold)
    return p


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    lines = text.split("\n")
    in_table = False
    in_code = False
    table_rows: list[list[str]] = []

    for line in lines:
        # 代码块
        if line.strip().startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            add_cnipa_styled_paragraph(doc, line, size=10, space_after=Pt(2))
            continue

        # 表格
        if line.strip().startswith("|") and line.strip().endswith("|"):
            if not in_table:
                in_table = True
                table_rows = []
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            # 跳过分隔行
            if all(re.match(r"^-+$", c) for c in cells):
                continue
            table_rows.append(cells)
            continue
        else:
            if in_table and table_rows:
                in_table = False
                _render_table(doc, table_rows)
                table_rows = []
            in_table = False

        # 空行
        if not line.strip():
            add_cnipa_styled_paragraph(doc, "", space_after=Pt(2))
            continue

        # 分割线
        if line.strip() == "---":
            add_cnipa_styled_paragraph(doc, "─" * 50, size=10, space_after=Pt(6))
            continue

        # 标题
        heading_match = re.match(r"^(#{1,4})\s+(.+)", line)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2)
            add_heading_styled(doc, title, level)
            continue

        # 普通段落
        cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        cleaned = re.sub(r"\*(.+?)\*", r"\1", cleaned)
        cleaned = re.sub(r"`(.+?)`", r"\1", cleaned)
        add_cnipa_styled_paragraph(doc, cleaned)

    # 处理最后的表格
    if in_table and table_rows:
        _render_table(doc, table_rows)

    doc.save(str(docx_path))


def _render_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols, style="Table Grid")
    for i, row_cells in enumerate(rows):
        for j, cell_text in enumerate(row_cells):
            if j >= max_cols:
                break
            cell = table.cell(i, j)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, CNIPA_FONT, CNIPA_FONT_EN, 10, False)
    doc.add_paragraph()


def main() -> None:
    if len(sys.argv) < 3:
        print(f"用法: {sys.argv[0]} <input.md> <output.docx>", file=sys.stderr)
        sys.exit(1)

    md_path = Path(sys.argv[1])
    docx_path = Path(sys.argv[2])

    if not md_path.is_file():
        print(f"错误: 输入文件不存在: {md_path}", file=sys.stderr)
        sys.exit(1)

    try:
        md_to_docx(md_path, docx_path)
        print(f"导出成功: {docx_path}")
    except Exception as e:
        print(f"导出失败: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
